# -*- coding: utf-8 -*-
# Part of Odoo. See LICENSE file for full copyright and licensing details.

import logging
import time
import sys
from odoo import api, fields, models
from odoo import tools, _
from odoo.exceptions import ValidationError
_logger = logging.getLogger(__name__)
from docx import Document
import xlsxwriter
class product_kks(models.Model):

    _name = "product.kks"
    _description = "KKS"
    _order = 'name'

    name = fields.Char('KKS')
    code = fields.Integer('Code KKS')
    item = fields.Integer('Item')
    reference = fields.Char('Référence')
    order_id = fields.Many2one('product.order','Commande')
    customer_id = fields.Many2one('res.partner','Client',domain=[('customer','=',True)],required=True)
    appareil_id = fields.Many2one('product.template','Appareil',domain=[('type','=','appareil')])
    maker_id = fields.Many2one('product.template.maker','Marque')
    ss_type_appareil_id = fields.Many2one('product.category','Sous Type Appareil',domain=[('type','=','ss_type')] )
    implantation_id = fields.Many2one('product.implantation', 'Implantation')
    type_implantation_id = fields.Many2one('product.type.implantation', 'Type Implantation')
    tarif = fields.Float('Prix m(DH)')
    arret_ids = fields.One2many('product.kks.arret','kks_id','Arrêts')
    tarif_ids = fields.One2many('product.kks.tarif','kks_id','Tarifs')
    echafaudage_ids = fields.One2many('product.kks.echafaudage','kks_id','Echafaudages') 
    piece_ids = fields.One2many('product.kks.piece','kks_id','Pièces') 
    pps_ids = fields.One2many('product.kks.pps','kks_id','PPS') 
    piece_prevoir_ = fields.Boolean('A Prévoir')
    piece_changee_ = fields.Boolean('Changée')
    choice_ = fields.Boolean('Choix')
    ############# add repere from echafaudage
    repere = fields.Float('Métrage(m3)')
    
    @api.onchange('piece_prevoir_')
    def _onchange_piece_prevoir_(self):
        for obj in self.piece_ids:
            obj.piece_prevoir=self.piece_prevoir_
    
    @api.onchange('piece_changee_')
    def _onchange_piece_changee_(self):
        for obj in self.piece_ids:
            obj.piece_changee=self.piece_changee_
        
    @api.onchange('choice_')
    def _onchange_choice_(self):
        for obj in self.piece_ids:
            obj.choice=self.choice_
            
    def open_unite(self):
        if self.piece_id:
            return {
                'type': 'ir.actions.act_window',
                'res_model': 'product.unite',
                'view_mode': 'tree',
                'view_type': 'form',
                'views': [(self.env.ref('arfi.product_unite_tree_view').id, 'tree')],
                }
    def write_file(self,data, filename):
        with open(filename, 'wb') as f:
            f.write(data)   
        
    def download_file(self,file_id):
        import base64
        reload(sys)
        sys.setdefaultencoding("utf-8")
        self._cr.execute("select filename,file_extension,type_file,file from muk_dms_file a\
                    left join muk_dms_database_data  b\
                    on a.file_ref=concat('muk_dms.database_data,',b.id)\
                    where a.id="+str(file_id))
        for res in self.env.cr.fetchall():
            filename=res[0]
            file_extension=res[1]
            data=res[3]
        destination='C:\\Users\\Yassine\\workspace\\arfi_2010\\addons\\web\\static\\reporting\\'+filename
        data=base64.decodestring(data)
        self.write_file(data, destination)
        if file_extension==".docx":               
            document = Document(destination)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if 'KKS' in paragraph.text:
                                paragraph.text=''
                                run = paragraph.add_run('KKS : '+self.name)
                                run.bold = True
            document.save(destination)
        url="/web/static/reporting/"+filename
        return {
            "type": "ir.actions.act_url",
            "url": url,
            "target": "new",
        }
    def action_print_pps(self):
        for obj in self.pps_ids:
            self.download_file(obj.file_id.id)
    def return_directory_id(self,name):
        rslt=None
        results = self.env['muk_dms.directory'].search([('name', '=', name)])
        for obj in results:
            rslt=obj.id
        return rslt
    def action_select_pps(self):
        self._cr.execute("update muk_dms_file set variant='t'")
        self._cr.commit()
        for obj in self.pps_ids:
            self._cr.execute("update muk_dms_file set variant='f' where id="+str(obj.file_id.id))
            self._cr.commit()
        return  {
                'type': 'ir.actions.act_window',
                'res_model': 'product.kks.pps.select',
                'view_mode': 'form',
                'view_type': 'form',
                'views': [(self.env.ref('arfi.view_product_kks_pps_select_form').id, 'form')],
                'context' : {'default_kks_id' : self.id,
                         'default_directory_id' : self.return_directory_id("Plans Preventions"),
                         'default_type_file' : 'prev'},
                'target': 'new',
                 }
    def action_select_tarif(self):
        self._cr.execute("update product_nature set variant='t'")
        self._cr.commit()
        for obj in self.tarif_ids:
            self._cr.execute("update product_nature set variant='f' where id="+str(obj.nature_id.id))
            self._cr.commit()
        return  {
                'type': 'ir.actions.act_window',
                'res_model': 'product.kks.tarif.wizard',
                'view_mode': 'form',
                'view_type': 'form',
                'views': [(self.env.ref('arfi.view_product_kks_tarif_wizard_form').id, 'form')],
                'context' : {'default_kks_id' : self.id},
                'target': 'new',
                 }
    def action_select_piece(self):
        self._cr.execute("update product_piece set variant='f'")
        self._cr.commit()
        self._cr.execute("update product_piece set variant='t' where appareil_id="+str(self.appareil_id.id))
        self._cr.commit()
        """for obj in self.piece_ids:
            self._cr.execute("update product_piece set variant='f' where id="+str(obj.piece_id.id))
            self._cr.commit()"""
        return  {
                'type': 'ir.actions.act_window',
                'res_model': 'product.kks.piece.wizard',
                'view_mode': 'form',
                'view_type': 'form',
                'views': [(self.env.ref('arfi.view_product_kks_piece_wizard_form').id, 'form')],
                'context' : {'default_kks_id' : self.id,
                              'default_appareil_id' : self.appareil_id.id},
                'target': 'new',
                 }
    @api.onchange('appareil_id')
    def _onchange_appareil(self):
        self.ss_type_appareil_id=self.appareil_id.ss_categ_id.id
        self.maker_id=self.appareil_id.maker_id.id
        self.reference=self.appareil_id.name
        
    @api.multi
    def print_product_kks_echafaudage(self):
        data = {}
        data['form'] = self.read(['id'])[0]
        return self._print_product_kks_echafaudage(data)
    def _print_product_kks_echafaudage(self, data):
        data['form'].update(self.read(['id'])[0])
        print data['form']
        return self.env['report'].get_action(self, 'arfi.action_report_productkksechafaudage', data=data)
    
    def action_observation(self):
        import sys
        reload(sys)
        sys.setdefaultencoding("utf-8")
        text1='<p width="100%" style="font-size:13px;">NB: Piéces de rechange à prévoir Pour Révision : </p>'
        text2='<br/><br/><p width="100%" style="font-size:13px;">NB: Piéces de rechange Changée Pour Révision : </p>'
        for record in self.piece_ids:
            magasin=ref_fab=piece=no_piece=''
            if record.magasin_id.code:
                magasin=record.magasin_id.code
            if record.piece_id.no_piece:
                no_piece=record.piece_id.no_piece
            if  record.piece_id.name:
                piece=record.piece_id.name
            if record.ref_fab:
                ref_fab=record.ref_fab
            if record.piece_prevoir==True:
                text1+='<span width="100%" style="font-size:13px;"><span style="width:15%;display: inline-block;"> -Rep : '+no_piece.replace("'", "''")+' </span><span style="width:40%;display: inline-block;"> '+piece.replace("'", "''")+' </span><span style="width:15%;display: inline-block;"> Réf : '+ref_fab.replace("'", "''")+' </span><span style="display: inline-block;"> code : '+magasin.replace("'", "''")+' </span></span></br>'
            if record.piece_changee==True:
                text2+='<span width="100%" style="font-size:13px;"><span style="width:15%;display: inline-block;"> -Rep : '+no_piece.replace("'", "''")+' </span><span style="width:40%;display: inline-block;"> '+piece.replace("'", "''")+' </span><span style="width:15%;display: inline-block;"> Réf : '+ref_fab.replace("'", "''")+' </span><span style="display: inline-block;"> code : '+magasin.replace("'", "''")+' </span></span></br>'
            order_id=record.kks_id.order_id.id
        text1+=''
        text2+=''
        if order_id:
            self._cr.execute("update product_order set obs_devis='"+text2+"',obs_recap='"+text2+"',obs_atelier='"+text1+""+text2+"'\
                                where id="+str(order_id))
            return {
                'type': 'ir.actions.act_window',
                'res_model': 'product.order',
                'view_mode': 'form',
                'view_type': 'form',
                'res_id': order_id,
                'views': [(self.env.ref('arfi.product_order_observation_create_view_form').id, 'form')],
                'target': 'new',
                'flags': {'form': {'action_buttons': True}}
                 }
        else :
            raise ValidationError(_("LE KKS n'est attaché à une commande"))
    def action_observation_(self):
        import sys
        reload(sys)
        sys.setdefaultencoding("utf-8")
        text1='<table width="100%" style="font-size:13px;margin-top: -30px;"><tr><td colspan="4">NB: Piéces de rechange à prévoir Pour Révision : </td></tr><tr><td colspan="4"></td></tr>'
        text2='<table width="100%" style="font-size:13px;"><tr><td colspan="4">NB: Piéces de rechange Changée Pour Révision : </td></tr><tr><td colspan="4"></td></tr>'
        for record in self.piece_ids:
            magasin=ref_fab=piece=no_piece=''
            if record.magasin_id.code:
                magasin=record.magasin_id.code
            if record.piece_id.no_piece:
                no_piece=record.piece_id.no_piece
            if  record.piece_id.name:
                piece=record.piece_id.name
            if record.ref_fab:
                ref_fab=record.ref_fab
            if record.piece_prevoir==True:
                text1+="<tr><td> -Rep : "+no_piece.replace("'", "''")+" </td><td > "+piece.replace("'", "''")+" </td><td > Réf : "+ref_fab.replace("'", "''")+" </td><td > code : "+magasin.replace("'", "''")+" </td></tr></br> "
            if record.piece_changee==True:
                text2+="<tr><td> -Rep : "+no_piece.replace("'", "''")+" </td><td > "+piece.replace("'", "''")+" </td><td > Ref : "+ref_fab.replace("'", "''")+" </td><td > code : "+magasin.replace("'", "''")+" </td></tr></br> "
            order_id=record.kks_id.order_id.id
        text1+='</table>'
        text2+='</table>'
        if order_id:
            self._cr.execute("update product_order set obs_devis='"+text2+"',obs_recap='"+text2+"',obs_atelier='"+text1+""+text2+"'\
                                where id="+str(order_id))
            return {
                'type': 'ir.actions.act_window',
                'res_model': 'product.order',
                'view_mode': 'form',
                'view_type': 'form',
                'res_id': order_id,
                'views': [(self.env.ref('arfi.product_order_observation_create_view_form').id, 'form')],
                'target': 'new',
                'flags': {'form': {'action_buttons': True}}
                 }
        else :
            raise ValidationError(_("LE KKS n'est attaché à une commande"))
    def action_pps(self):
   
        return {
            'type': 'ir.actions.act_window',
            'res_model': 'product.kks',
            'view_mode': 'form',
            'view_type': 'form',
            'res_id': self.id,
            'views': [(self.env.ref('arfi.product_kks_risque_create_view_form').id, 'form')],
            #'target': 'new',
            'flags': {'form': {'action_buttons': True}}
             }   
    def action_echafaudage(self):
   
        return {
            'type': 'ir.actions.act_window',
            'res_model': 'product.kks',
            'view_mode': 'form',
            'view_type': 'form',
            'res_id': self.id,
            'views': [(self.env.ref('arfi.product_kks_echafaudage_create_view_form').id, 'form')],
            'target': 'new',
            'flags': {'form': {'action_buttons': True}}
             }
    def action_pdr(self):
   
        return {
            'type': 'ir.actions.act_window',
            'res_model': 'product.kks',
            'view_mode': 'form',
            'view_type': 'form',
            'res_id': self.id,
            'views': [(self.env.ref('arfi.product_kks_pdr_create_view_form').id, 'form')],
            'target': 'self',
            'flags': {'form': {'action_buttons': True}}}

    @api.multi
    def action_pid(self):
        if self:
            return self.env['product.kks.pid.annotation'].action_open_annotation(self.name)
        raise ValidationError(_("Attention! Cette commande n'est affecté à aucun KKS."))
class product_kks_arret(models.Model):

    _name = "product.kks.arret"
    _description = "KKS Arrets"

    kks_id = fields.Many2one('product.kks','Code KKS')
    unite_id = fields.Many2one('product.unite','Code Unité')
    travaux_id = fields.Many2one('product.travaux','Travaux')
    arret_id = fields.Many2one('product.arret','Code Arrêt')
class product_kks_tarif(models.Model):

    _name = "product.kks.tarif"
    _description = "KKS Tarifs"

    kks_id = fields.Many2one('product.kks','Code KKS')
    customer_id = fields.Many2one('res.partner','Client',related='kks_id.customer_id',store=True, readonly=False)
    appareil_id = fields.Many2one('product.template','Appareil',related='kks_id.appareil_id',store=True, readonly=False)
    item = fields.Integer(related='kks_id.item',string='Item', store=True ,readonly=False)
    nature_id = fields.Many2one('product.nature','Nature Travaux')
    montant = fields.Float('Prix de vente')
    montant2 = fields.Float("Prix d'achat")
    choice = fields.Boolean('Choix')
    fact = fields.Selection([('Contrat','Contrat'),('BC','BC')],'Fact')
    
class product_kks_echafaudage(models.Model):

    _name = "product.kks.echafaudage"
    _description = "KKS Echafaudage"

    kks_id = fields.Many2one('product.kks','Code KKS')
    echafaudage_id = fields.Many2one('product.echafaudage','Désignation')
    qte = fields.Integer('Qte utilisée)')
class product_kks_piece(models.Model):

    _name = "product.kks.piece"
    _description = "KKS Piece"
    
    
    kks_id = fields.Many2one('product.kks','Code KKS')
    appareil_id = fields.Many2one('product.template','Appareil',domain=[('type','=','appareil')])
    choice = fields.Boolean('Choix')
    tarif = fields.Float('Tarif')
    magasin_id = fields.Many2one('product.magasin', 'Code Magasin')
    piece_id = fields.Many2one('product.piece', 'Désignation',domain="[('appareil_id','=',appareil_id),('appareil_id','!=',None')]",
                                change_default=True,)
    no_piece = fields.Char(related='piece_id.no_piece',string="REP", store=True ,readonly=True)
    ref_fab = fields.Char('Réf fabriquant')
    ref_com = fields.Char('Réf Commercial')
    datetarif_id = fields.Many2one('product.datetarif','Date Tarif')
    ########## Copier et coller
    piece_prevoir = fields.Boolean('A Prévoir')
    piece_changee = fields.Boolean('Changée')
    
    @api.multi
    def bulk_verify(self):
        text='<table>'
        
        for record in self:
            magasin=ref_fab=piece=no_piece=''
            if record.magasin_id.code:
                magasin=record.magasin_id.code
            if record.piece_id.no_piece:
                no_piece=record.piece_id.no_piece
            if  record.piece_id.name:
                piece=record.piece_id.name
            if record.ref_fab:
                ref_fab=record.ref_fab
            text+='<tr><td>-Rep '+no_piece+'</td><td>'+piece+'</td><td>'+ref_fab+'</td><td>code '+magasin+"</td></br>"
            order_id=record.kks_id.order_id.id
        text+='</table>'
        if order_id:
            self._cr.execute("update product_order set obs_devis='"+text+"',obs_recap='"+text+"',obs_atelier='"+text+"'\
                                where id="+str(order_id))
        else :
            raise ValidationError(_("LE KKS n'est attaché à une commande"))
    def open_magasin(self):
        if self.magasin_id:
            return {
                'type': 'ir.actions.act_window',
                'res_model': 'product.magasin',
                'view_mode': 'form',
                'view_type': 'form',
                'res_id': self.magasin_id.id,
                'views': [(self.env.ref('arfi.product_magasin_form_view').id, 'form')],
                'target': 'new',
                'flags': {'form': {'action_buttons': True}}}
    def open_piece(self):
        if self.piece_id:
            return {
                'type': 'ir.actions.act_window',
                'res_model': 'product.magasin',
                'view_mode': 'form',
                'view_type': 'form',
                'res_id': self.magasin_id.id,
                'views': [(self.env.ref('arfi.product_magasin_photo_form_view').id, 'form')],
                'target': 'new',
                'flags': {'form': {'action_buttons': True}}}
    def action_magasin(self):
   
        return {
            'type': 'ir.actions.act_window',
            'res_model': 'product.magasin',
            'view_mode': 'form',
            'view_type': 'form',
            'res_id': self.magasin_id.id,
            'views': [(self.env.ref('arfi.product_magasin_form_view').id, 'form')],
            'target': 'new',
            'flags': {'form': {'action_buttons': True}}}
class product_kks_supplier(models.Model):

    _name = "product.kks.supplier"
    _description = "KKS Fournisseurs"
    
    code = fields.Char('Code')
    supplier_id = fields.Many2one('res.supplier','Fournisseur')
    contact_id = fields.Many2one('res.contact','Contact')
    magasin_id = fields.Many2one('product.magasin','Magasin')
    adress = fields.Char('Adresse')
    gsm = fields.Char('GSM')
    tel = fields.Char('TEL')
    fax = fields.Char('FAX')
    email = fields.Char('Email')
    price = fields.Float('Prix')
    date = fields.Date('Date')
    
    @api.onchange('contact_id') # if these fields are changed, call method
    def onchange_contact_id(self):
        self.tel = self.contact_id.phone
        self.gsm = self.contact_id.mobile
        self.adress = self.contact_id.street
        self.fax = self.contact_id.fax
        self.email = self.contact_id.email
class product_kks_stock(models.Model):

    _name = "product.kks.stock"
    _description = "KKS Stock"
    _order = 'sequence'

    magasin_id = fields.Many2one('product.magasin','Magasin')
    info_id = fields.Many2one('product.info','Info')
    value = fields.Char('Valeur')
    sequence = fields.Integer(related='info_id.sequence', string="Sequence", store=True, readonly=True)
    
class product_magasin(models.Model):
    _name = 'product.magasin'
    _rec_name = 'code'
    _order = 'code'
    
    def action_select_info(self):
        self._cr.execute("update product_info set variant='t'")
        self._cr.commit()
        for obj in self.stock_ids:
            self._cr.execute("update product_info set variant='f' where id="+str(obj.info_id.id))
            self._cr.commit()
        return  {
                'type': 'ir.actions.act_window',
                'res_model': 'product.magasin.stock.wizard',
                'view_mode': 'form',
                'view_type': 'form',
                'views': [(self.env.ref('arfi.view_product_magasin_stock_wizard_form').id, 'form')],
                'context' : {'default_magasin_id' : self.id},
                'target': 'new',
                 }
        
    code = fields.Char('Code', required=True)
    photo_name = fields.Char('Nom du fichier',size=256)
    photo = fields.Binary("Image")
    supplier_ids =  fields.One2many('product.kks.supplier','magasin_id',string='Fournisseurs', readonly=False)
    stock_ids =  fields.One2many('product.kks.stock','magasin_id',string='Stock', readonly=False)
    purchase_ids =  fields.One2many('product.purchase.line','magasin_id',string='Achats', readonly=True)
    piece_ids =  fields.One2many('product.kks.piece','magasin_id',string='Pieces', readonly=True)
    
class product_kks_facture(models.Model):

    _name = "product.kks.facture"
    _description = "Product KKs Facture"

    @api.depends('facture_line','facture_line')
    def _amount_all(self):
        """
        Compute the total amounts of the SO.
        """
        for order in self:
            montant_ht = 0.0
            for line in order.facture_line:
                montant_ht+=line.montant
            order.update({
                'montant_ht': montant_ht,
                'tva': montant_ht*0.2,
                'montant_ttc': montant_ht*1.2,
            })
            
    name = fields.Char('Facture')
    type_extraction = fields.Char('Type')
    date_facture = fields.Date('Date facture' ,default=time.strftime('%Y-%m-%d'))
    no_facture = fields.Char('Facture n°')
    type_facture = fields.Selection([('DEM','DEM'),('SMP','SMP')], 'Type facture',default='DEM')
    no_contrat = fields.Char('BC/Contract n°')
    type = fields.Selection([('facture','Facture'),('devis','Devis')], 'Type')
    no_devis = fields.Char('Devis n°')
    no_reference = fields.Char('Référence n°')
    unite_id = fields.Many2one('product.unite','Unité',readonly=True)
    customer_id = fields.Many2one('res.partner','Client',readonly=True)
    arret_id = fields.Many2one('product.arret','Code Arrêt',readonly=True)
    facture_line = fields.One2many('product.kks.facture.line','facture_id','Lignes', readonly=True)
    montant_ht = fields.Float(compute='_amount_all',string='Montant HT', readonly=True, store=True)
    tva = fields.Float(compute='_amount_all',string='TVA (20%)', readonly=True, store=True)
    montant_ttc = fields.Float(compute='_amount_all',string='Montant TTC', readonly=True, store=True)
    montant_text = fields.Text('Montant', readonly=True)
    
    @api.multi
    def print_facture(self):
        data = {}
        data['form'] = self.read(['id'])[0]
        return self._print_facture(data)
    def _print_facture(self, data):
        data['form'].update(self.read(['id'])[0])
        return self.env['report'].get_action(self, 'arfi.action_report_productkksfacturechangement', data=data)
    @api.multi
    def print_pv(self):
        data = {}
        data['form'] = self.read(['id'])[0]
        return self._print_pv(data)
    def _print_pv(self, data):
        data['form'].update(self.read(['id'])[0])
        return self.env['report'].get_action(self, 'arfi.action_report_productkkspvchangement', data=data)
    
    def get_return(self,fichier):
        url = '/web/static/reporting/' + fichier
        if url:
            return { 'type': 'ir.actions.act_url','target': 'new', 'url': url, 'nodestroy': True }
        else:
            return True
    @api.multi
    def print_pv_excel(self):
        data = {}
        data['form'] = self.read(['id'])[0]
        return self._print_pv_excel(data)
    def _print_pv_excel(self,data):
        reload(sys)
        sys.setdefaultencoding("UTF8")
        results = self.env['product.image.directory'].search([('type','=','reporting')])
        for result in results:
            directory=result.name
        fichier="PV _"+time.strftime('%H%M%S')+".xlsx"
        workbook = xlsxwriter.Workbook(directory + fichier)
        style_title = workbook.add_format({'bg_color' : '#003366','color' : 'white','text_wrap' : True,'bold' :1,  'align' : 'center','valign' : 'vcenter',
                                                          'top' : 1, 'bottom' : 1,})
        style_titre = workbook.add_format({'color' : '#003366','text_wrap' : True,'bold' :1,  'align' : 'center','valign' : 'vcenter',
                                                          'font_size' : 18, 'font_name' : 'tahoma'})
        style_titre2 = workbook.add_format({'text_wrap' : True,'bold' :1,  'align' : 'center','valign' : 'vcenter',
                                                          'font_size' : 14, 'font_name' : 'tahoma' ,'border' : 2})
        style = workbook.add_format({'text_wrap' : True,'bold' :1,  'align' : 'center','valign' : 'vcenter','top' : 1, 'bottom' : 1, })
        style_ = workbook.add_format({'text_wrap' : True,'bold' :1,   'top' : 1, 'bottom' : 1,})
        style1 = workbook.add_format({  'text_wrap' : True,  'align' : 'center','valign' : 'vcenter','top' : 1, 'bottom' : 1,})
        style1_ = workbook.add_format({  'text_wrap' : True,  'top' : 1, 'bottom' : 1,})
        feuille=workbook.add_worksheet('PDR')
        feuille.set_zoom(85)
        feuille.freeze_panes(5, 0)
        feuille.set_tab_color('yellow')
        feuille.set_column('A:A', 6)
        feuille.set_column('B:B', 15)
        feuille.set_column('C:C', 53)
        feuille.set_column('D:D', 22)
        feuille.set_column('E:E', 22)
        feuille.merge_range('A1:E1', 'PV', style_titre)
        record = self.env['product.kks.facture'].search([],limit=1)
        titre='Désignation '+record.type_extraction+' arrêt unité '+record.unite_id.code+' du '+record.arret_id.date
        feuille.merge_range('A3:E3', titre, style_titre2)
        x=5
        feuille.write('A'+str(x),'item',style_title)
        feuille.write('B'+str(x),'KKS',style_title)
        feuille.write('C'+str(x),'Réference',style_title)
        feuille.write('D'+str(x),'Marque',style_title)
        feuille.write('E'+str(x),'Nature Travaux',style_title)
        
        records = self.env['product.kks.facture.line'].search([])
        x=x+1
        for record in records:         
            feuille.write('A'+str(x),record.item,style)
            feuille.write('B'+str(x),record.kks,style)
            feuille.write('C'+str(x),record.reference,style_)
            feuille.write('D'+str(x),record.maker_id.name,style)
            feuille.write('E'+str(x),record.nature_id.name,style)
            x=x+1

        workbook.close()
        return self.get_return(fichier)
    @api.multi
    def print_facture_excel(self):
        data = {}
        data['form'] = self.read(['id'])[0]
        return self._print_facture_excel(data)
    def _print_facture_excel(self,data):
        reload(sys)
        sys.setdefaultencoding("UTF8")
        results = self.env['product.image.directory'].search([('type','=','reporting')])
        for result in results:
            directory=result.name
        fichier="Facture _"+time.strftime('%H%M%S')+".xlsx"
        workbook = xlsxwriter.Workbook(directory + fichier)
        style_title = workbook.add_format({'bg_color' : '#003366','color' : 'white','text_wrap' : True,'bold' :1,  'align' : 'center','valign' : 'vcenter',
                                                          'top' : 1, 'bottom' : 1,})
        style_titre = workbook.add_format({'color' : '#003366','text_wrap' : True,'bold' :1,  'align' : 'center','valign' : 'vcenter',
                                                          'font_size' : 18, 'font_name' : 'tahoma'})
        style_titre2 = workbook.add_format({'text_wrap' : True,'bold' :1,  'align' : 'center','valign' : 'vcenter',
                                                          'font_size' : 14, 'font_name' : 'tahoma' ,'border' : 2})
        style = workbook.add_format({'text_wrap' : True,'bold' :1,  'align' : 'center','valign' : 'vcenter','top' : 1, 'bottom' : 1, })
        style_nombre = workbook.add_format({'num_format' : '#,##0.00','text_wrap' : True,'bold' :1,  'align' : 'center','valign' : 'vcenter','top' : 1, 'bottom' : 1, })
        style_ = workbook.add_format({'text_wrap' : True,'bold' :1,   'top' : 1, 'bottom' : 1,})
        style1 = workbook.add_format({  'text_wrap' : True,  'align' : 'center','valign' : 'vcenter','top' : 1, 'bottom' : 1,})
        style1_ = workbook.add_format({  'text_wrap' : True,  'top' : 1, 'bottom' : 1,})
        feuille=workbook.add_worksheet('PDR')
        feuille.set_zoom(85)
        feuille.freeze_panes(5, 0)
        feuille.set_tab_color('yellow')
        feuille.set_column('A:A', 6)
        feuille.set_column('B:B', 15)
        feuille.set_column('C:C', 53)
        feuille.set_column('D:F', 22)
        feuille.merge_range('A1:F1', 'Facture', style_titre)
        record = self.env['product.kks.facture'].search([],limit=1)
        titre='Désignation '+record.type_extraction+' arrêt unité '+record.unite_id.code+' du '+record.arret_id.date
        feuille.merge_range('A3:F3', titre, style_titre2)
        x=5
        feuille.write('A'+str(x),'item',style_title)
        feuille.write('B'+str(x),'KKS',style_title)
        feuille.write('C'+str(x),'Réference',style_title)
        feuille.write('D'+str(x),'Marque',style_title)
        feuille.write('E'+str(x),'Nature Travaux',style_title)
        feuille.write('F'+str(x),'Montant',style_title)
        
        records = self.env['product.kks.facture.line'].search([])
        x=x+1
        for record in records:         
            feuille.write('A'+str(x),record.item,style)
            feuille.write('B'+str(x),record.kks,style)
            feuille.write('C'+str(x),record.reference,style_)
            feuille.write('D'+str(x),record.maker_id.name,style)
            feuille.write('E'+str(x),record.nature_id.name,style)
            feuille.write('F'+str(x),record.montant,style_nombre)
            x=x+1

        workbook.close()
        return self.get_return(fichier)
    
    @api.multi
    def print_facture_changement(self):
        data = {}
        data['form'] = self.read(['id'])[0]
        return self._print_facture_changement(data)
    def _print_facture_changement(self, data):
        data['form'].update(self.read(['id'])[0])
        return self.env['report'].get_action(self, 'arfi.action_report_productkksfacturechangement', data=data)
    @api.multi
    def print_pv_changement(self):
        data = {}
        data['form'] = self.read(['id'])[0]
        return self._print_pv_changement(data)
    def _print_pv_changement(self, data):
        data['form'].update(self.read(['id'])[0])
        return self.env['report'].get_action(self, 'arfi.action_report_productkkspvchangement', data=data)
class product_kks_facture_line(models.Model):

    _name = "product.kks.facture.line"
    _description = "Product KKs Facture Line"
    _order ='kks'

    reference = fields.Char('Référence')
    facture_id = fields.Many2one('product.kks.facture', 'Facture')
    kks_id = fields.Many2one('product.kks','KKS')
    kks = fields.Char(related='kks_id.name',string="KKS", store=True ,readonly=True)
    item = fields.Integer('Item')
    maker_id = fields.Many2one('product.template.maker','Marque')
    unite_id = fields.Many2one('product.unite','Code Unité')
    customer_id = fields.Many2one('res.partner','Client')
    arret_id = fields.Many2one('product.arret','Code Arrêt')
    montant = fields.Float('Montant (Dhs)')
    nature_id = fields.Many2one('product.nature','Nature Travaux')
    travaux_id = fields.Many2one('product.travaux','Travaux')
    fact = fields.Selection([('Contrat','Contrat'),('BC','BC')],'Fact')
class product_kks_pps(models.Model):
    _name = 'product.kks.pps'

    kks_id = fields.Many2one('product.kks','KKS')
    directory_id = fields.Many2one('muk_dms.directory', string="Directory", required=False)
    file_id = fields.Many2one('muk_dms.file', 'Fichier',
                                domain="[('directory','=',directory_id)]",
                                change_default=True,help="")
    type_file = fields.Selection([('prev','Plan Prévention')] ,'Type du fichier' , default='prev', required=False)
    print_file = fields.Boolean('A Imprimer?')
    
    def write_file(self,data, filename):
        with open(filename, 'wb') as f:
            f.write(data)
    
        
    def download_file(self):
        import base64
        reload(sys)
        sys.setdefaultencoding("utf-8")
        results = self.env['product.image.directory'].search([('type','=','reporting')])
        for result in results:
            directory=result.name
        self._cr.execute("select filename,file_extension,type_file,file from muk_dms_file a\
                    left join muk_dms_database_data  b\
                    on a.file_ref=concat('muk_dms.database_data,',b.id)\
                    where a.id="+str(self.file_id.id))
        for res in self.env.cr.fetchall():
            filename=res[0]
            file_extension=res[1]
            data=res[3]
        filename=filename.split('.')[0]+'_'+time.strftime('%H%M%S')+file_extension
        #destination='C:\\Users\\Yassine\\workspace\\arfi_2010\\addons\\web\\static\\reporting\\'+filename
        destination=directory+filename
        data=base64.decodestring(data)
        self.write_file(data, destination)
        if file_extension==".docx": 
            document = Document(destination)    
            from docx.shared import Pt
            style = document.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(10)                
            
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if 'KKS' in paragraph.text:
                                paragraph.text=''
                                run = paragraph.add_run('KKS : '+self.kks_id.name)
                                run.bold = True
                                paragraph.style = document.styles['Normal']
                            if 'Client' in paragraph.text:
                                paragraph.text=''
                                run = paragraph.add_run('Client : '+self.kks_id.customer_id.name)
                                run.bold = True
                                paragraph.style = document.styles['Normal']
                            
            document.save(destination)
        url="/web/static/reporting/"+filename
        return {
            "type": "ir.actions.act_url",
            "url": url,
            "target": "new",
        }
    def download_file__(self):
        import base64
        reload(sys)
        sys.setdefaultencoding("utf-8")
        self._cr.execute("select filename,file_extension,type_file,file from muk_dms_file a\
                    left join muk_dms_database_data  b\
                    on a.file_ref=concat('muk_dms.database_data,',b.id)\
                    where a.id="+str(self.file_id.id))
        for res in self.env.cr.fetchall():
            filename=res[0]
            file_extension=res[1]
            data=res[3]
       
        if file_extension==".docx":  
            destination='/usr/lib/python2.7/dist-packages/odoo/addons/web/static/reporting/'+filename
            ##destination='C:\\Users\\Yassine\\workspace\\arfi_2010\\addons\\web\\static\\reporting\\'+filename
            data=base64.decodestring(data)
            self.write_file(data, destination)
                  
            document = Document(destination)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if 'KKS' in paragraph.text:
                                paragraph.text=''
                                run = paragraph.add_run('KKS : '+self.kks_id.name)
                                run.bold = True
            document.save(destination)
            url="/web/static/reporting/"+filename
        else :
            url="/dms/file/download/"+str(self.file_id.id)
        return {
            "type": "ir.actions.act_url",
            "url": url,
            "target": "new",
        }
    def return_directory_id(self,name):
        rslt=None
        results = self.env['muk_dms.directory'].search([('name', '=', name)])
        for obj in results:
            rslt=obj.id
        return rslt
    @api.onchange('type_file')
    def _onchange_type_file(self):
        if self.type_file=="prev":
            self.directory_id=self.return_directory_id("Plans Preventions")
        else :
            self.directory_id=None