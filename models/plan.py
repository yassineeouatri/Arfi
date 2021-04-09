# -*- coding: utf-8 -*-

###################################################################################
# 
#    MuK Document Management System
#
#    Copyright (C) 2017 MuK IT GmbH
#
#    This program is free software: you can redistribute it and/or modify
#    it under the terms of the GNU Affero General Public License as
#    published by the Free Software Foundation, either version 3 of the
#    License, or (at your option) any later version.
#
#    This program is distributed in the hope that it will be useful,
#    but WITHOUT ANY WARRANTY; without even the implied warranty of
#    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
#    GNU Affero General Public License for more details.
#
#    You should have received a copy of the GNU Affero General Public License
#    along with this program.  If not, see <http://www.gnu.org/licenses/>.
#
###################################################################################

import os
import base64
import json
import urllib
import StringIO
import cStringIO
import mimetypes
import logging

from odoo import _
from odoo import models, api, fields
from odoo.exceptions import ValidationError, AccessError

from odoo.addons.muk_dms.models import muk_dms_base as base
from docx import Document
_logger = logging.getLogger(__name__)

class plan_prevention(models.Model):
    _name = 'plan.prevention'
    _description = "Plans Prevention"
    
        
    filename = fields.Char(string='Filename', required=True)
    file = fields.Binary(string='File', required=True)
    
    file_extension = fields.Char(string='Extension', readonly=True)
    mime_type = fields.Char(string='Type', readonly=True)
    file_size = fields.Integer(string='Size', readonly=True)
    
    name = fields.Char(compute='_compute_name', string="Name")
    

    
    
    #----------------------------------------------------------
    # Read, View 
    #----------------------------------------------------------

    @api.one
    @api.depends('filename')
    def _compute_name(self):
        self.name = self.filename and os.path.splitext(self.filename)[0] or ""

class plans(base.DMSModel):
    _inherit = 'muk_dms.file'
    _order = 'filename'
    
    variant = fields.Boolean('Variant',default=True)
    type_file = fields.Selection([('prev','Plan Prévention'),('ins','Instruction de travail'),
                             ('ope','Mode Opératoire'),('man','Manuel de Maintenance'),
                             ('orig','Plan Original'),('codif','Codification'),
                             ('spec', 'Spécification de Qualité'),
                             ('mos', 'Mode Opératoire de Soudage')
                             ] ,'Type du fichier' , required=True)
    
    
    def return_directory_id(self,name):
        rslt=None
        results = self.env['muk_dms.directory'].search([('name', '=', name)])
        for obj in results:
            rslt=obj.id
        return rslt
    def write_file(self,data, filename):
        with open(filename, 'wb') as f:
            f.write(data)
    def download_file(self):
        import base64
        import time
        import sys
        reload(sys)
        sys.setdefaultencoding("utf-8")
        results = self.env['product.image.directory'].search([('type','=','reporting')])
        for result in results:
            directory=result.name
        self._cr.execute("select filename,file_extension,type_file,file from muk_dms_file a\
                    left join muk_dms_database_data  b\
                    on a.file_ref=concat('muk_dms.database_data,',b.id)\
                    where a.id="+str(self.id))
        for res in self.env.cr.fetchall():
            filename=res[0]
            file_extension=res[1]
            data=res[3]
        filename=filename.split('.')[0]+'_'+time.strftime('%H%M%S')+file_extension
        #destination='C:\\Users\\Yassine\\workspace\\arfi_2010\\addons\\web\\static\\reporting\\'+filename
        destination=directory+filename
        data=base64.decodestring(data)
        self.write_file(data, destination)
        """if file_extension==".docx": 
            document = Document(destination)    
            from docx.shared import Pt
            style = document.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(10)                
            
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if 'KKS' in paragraph.text:
                                paragraph.text=''
                                run = paragraph.add_run('KKS : '+self.kks_id.name)
                                run.bold = True
                                paragraph.style = document.styles['Normal']
                            if 'Client' in paragraph.text:
                                paragraph.text=''
                                run = paragraph.add_run('Client : '+self.kks_id.customer_id.name)
                                run.bold = True
                                paragraph.style = document.styles['Normal']
                            
            document.save(destination)"""
        url="/web/static/reporting/"+filename
        return {
            "type": "ir.actions.act_url",
            "url": url,
            "target": "new",
        }
    @api.onchange('type_file')
    def _onchange_type_file(self):
        if self.type_file=="prev":
            self.directory=self.return_directory_id("Plans Preventions")
        elif self.type_file=="ins":
            self.directory=self.return_directory_id("Instructions de Travail")
        elif self.type_file=="ope":
            self.directory=self.return_directory_id("Mode Operatoire")
        elif self.type_file=="man":
            self.directory=self.return_directory_id("Manuel de Maintenance")
        elif self.type_file=="codif":
            self.directory=self.return_directory_id("Codification")
        elif self.type_file=="spec":
            self.directory=self.return_directory_id("Specification de qualite")
        elif self.type_file=="mos":
            self.directory=self.return_directory_id("Mode Operatoire de Soudage")
        else :
            self.directory=None